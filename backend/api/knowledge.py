#!/usr/bin/env python3
"""
知识库系统API
提供文档管理、知识沉淀、全文搜索功能
增强版：支持分类树、多级分类、文章状态、批量操作、权限控制
"""

from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from datetime import datetime
from sqlalchemy import or_, func, desc
from sqlalchemy.orm import joinedload, subqueryload
from enhanced_app import db, KnowledgeCategory, KnowledgeArticle, KnowledgeAttachment, KnowledgeComment, KnowledgeFavorite, KnowledgeReadRecord, User
import os
import uuid
import logging
import json

logger = logging.getLogger(__name__)
knowledge_bp = Blueprint('knowledge', __name__, url_prefix='/knowledge')


def check_admin():
    """检查是否为管理员"""
    try:
        current_user_id = get_jwt_identity()
        from enhanced_app import User
        user = User.query.get(current_user_id)
        if user and user.role == 'admin':
            return True
        return False
    except Exception:
        return False


def get_current_user():
    """获取当前用户"""
    try:
        current_user_id = get_jwt_identity()
        from enhanced_app import User
        return User.query.get(current_user_id)
    except Exception:
        return None


def can_edit_category(category):
    """检查用户是否有权限编辑分类"""
    if check_admin():
        return True
    current_user_id = get_jwt_identity()
    return category.created_by == current_user_id


def can_edit_article(article):
    """检查用户是否有权限编辑文章"""
    if check_admin():
        return True
    current_user_id = get_jwt_identity()
    return article.author_id == current_user_id


def get_all_child_category_ids(category_id):
    """递归获取所有子分类ID"""
    child_ids = [category_id]
    children = KnowledgeCategory.query.filter_by(parent_id=category_id).all()
    for child in children:
        child_ids.extend(get_all_child_category_ids(child.id))
    return child_ids


def update_category_article_count(category_id):
    """更新分类的文章数量（包括子分类）"""
    category = KnowledgeCategory.query.get(category_id)
    if not category:
        return

    child_ids = get_all_child_category_ids(category_id)
    count = KnowledgeArticle.query.filter(
        KnowledgeArticle.category_id.in_(child_ids),
        KnowledgeArticle.status == 'published'
    ).count()

    category.article_count = count
    db.session.commit()


# ==================== 分类管理 ====================
@knowledge_bp.route('/categories', methods=['GET'])
@jwt_required()
def get_categories():
    """获取分类列表（树状结构）"""
    try:
        include_archived = request.args.get('include_archived', 'false').lower() == 'true'
        flat = request.args.get('flat', 'false').lower() == 'true'

        if flat:
            query = KnowledgeCategory.query
            if not include_archived:
                query = query.filter_by(is_archived=0)
            categories = query.order_by(KnowledgeCategory.sort_order, KnowledgeCategory.name).all()
            return jsonify({
                'success': True,
                'data': [cat.to_dict(include_children=False) for cat in categories]
            })
        else:
            root_categories = KnowledgeCategory.query.filter_by(parent_id=None)
            if not include_archived:
                root_categories = root_categories.filter_by(is_archived=0)
            root_categories = root_categories.order_by(KnowledgeCategory.sort_order, KnowledgeCategory.name).all()

            def build_tree(cat):
                children = KnowledgeCategory.query.filter_by(parent_id=cat.id)
                if not include_archived:
                    children = children.filter_by(is_archived=0)
                children = children.order_by(KnowledgeCategory.sort_order, KnowledgeCategory.name).all()

                cat_dict = cat.to_dict(include_children=False)
                if children:
                    cat_dict['children'] = [build_tree(child) for child in children]
                else:
                    cat_dict['children'] = []
                return cat_dict

            tree = [build_tree(cat) for cat in root_categories]
            return jsonify({
                'success': True,
                'data': tree
            })
    except Exception as e:
        logger.error(f"获取分类错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/categories/tree', methods=['GET'])
@jwt_required()
def get_category_tree():
    """获取分类树（带文章数量统计）"""
    try:
        include_archived = request.args.get('include_archived', 'false').lower() == 'true'

        def get_article_count_for_category(cat_id):
            child_ids = get_all_child_category_ids(cat_id)
            query = KnowledgeArticle.query.filter(
                KnowledgeArticle.category_id.in_(child_ids),
                KnowledgeArticle.status == 'published'
            )
            return query.count()

        def build_tree_with_count(cat):
            children = KnowledgeCategory.query.filter_by(parent_id=cat.id)
            if not include_archived:
                children = children.filter_by(is_archived=0)
            children = children.order_by(KnowledgeCategory.sort_order, KnowledgeCategory.name).all()

            cat_dict = cat.to_dict(include_children=False, include_article_count=False)
            cat_dict['article_count'] = get_article_count_for_category(cat.id)

            if children:
                cat_dict['children'] = [build_tree_with_count(child) for child in children]
            else:
                cat_dict['children'] = []
            return cat_dict

        root_categories = KnowledgeCategory.query.filter_by(parent_id=None)
        if not include_archived:
            root_categories = root_categories.filter_by(is_archived=0)
        root_categories = root_categories.order_by(KnowledgeCategory.sort_order, KnowledgeCategory.name).all()

        tree = [build_tree_with_count(cat) for cat in root_categories]
        return jsonify({
            'success': True,
            'data': tree
        })
    except Exception as e:
        logger.error(f"获取分类树错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/categories', methods=['POST'])
@jwt_required()
def create_category():
    """创建分类"""
    try:
        data = request.get_json()
        current_user_id = get_jwt_identity()

        if not data.get('name', '').strip():
            return jsonify({'success': False, 'error': '分类名称不能为空'}), 400

        parent_id = data.get('parent_id')
        if parent_id:
            parent = KnowledgeCategory.query.get(parent_id)
            if not parent:
                return jsonify({'success': False, 'error': '父分类不存在'}), 400
            level = 1
            p = parent
            while p.parent_id:
                level += 1
                p = p.parent
            if level >= 3:
                return jsonify({'success': False, 'error': '分类层级不能超过3级'}), 400

        max_order = db.session.query(func.max(KnowledgeCategory.sort_order)).filter_by(parent_id=parent_id).scalar() or 0

        category = KnowledgeCategory(
            name=data['name'].strip(),
            description=data.get('description', ''),
            parent_id=parent_id,
            sort_order=data.get('sort_order', max_order + 1),
            created_by=current_user_id
        )
        db.session.add(category)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '分类创建成功',
            'data': {'id': category.id}
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"创建分类错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/categories/<int:cat_id>', methods=['GET'])
@jwt_required()
def get_category(cat_id):
    """获取单个分类详情"""
    try:
        category = KnowledgeCategory.query.get(cat_id)
        if not category:
            return jsonify({'success': False, 'error': '分类不存在'}), 404

        child_ids = get_all_child_category_ids(cat_id)
        article_count = KnowledgeArticle.query.filter(
            KnowledgeArticle.category_id.in_(child_ids),
            KnowledgeArticle.status == 'published'
        ).count()

        result = category.to_dict()
        result['article_count'] = article_count
        return jsonify({
            'success': True,
            'data': result
        })
    except Exception as e:
        logger.error(f"获取分类详情错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/categories/<int:cat_id>', methods=['PUT'])
@jwt_required()
def update_category(cat_id):
    """更新分类"""
    try:
        category = KnowledgeCategory.query.get(cat_id)
        if not category:
            return jsonify({'success': False, 'error': '分类不存在'}), 404

        if not can_edit_category(category):
            return jsonify({'success': False, 'error': '无权限编辑此分类'}), 403

        data = request.get_json()

        if 'name' in data:
            if not data['name'].strip():
                return jsonify({'success': False, 'error': '分类名称不能为空'}), 400
            category.name = data['name'].strip()
        if 'description' in data:
            category.description = data['description']
        if 'sort_order' in data:
            category.sort_order = data['sort_order']

        if 'parent_id' in data:
            new_parent_id = data['parent_id']
            if new_parent_id != category.id:
                if new_parent_id:
                    new_parent = KnowledgeCategory.query.get(new_parent_id)
                    if not new_parent:
                        return jsonify({'success': False, 'error': '父分类不存在'}), 400
                    level = 0
                    p = new_parent
                    while p.parent_id:
                        level += 1
                        p = p.parent
                    check_cat = category
                    while check_cat.parent_id:
                        level += 1
                        check_cat = check_cat.parent
                    if level >= 3:
                        return jsonify({'success': False, 'error': '移动后分类层级将超过3级'}), 400

                    child_ids = get_all_child_category_ids(category.id)
                    if new_parent_id in child_ids:
                        return jsonify({'success': False, 'error': '不能将分类移动到其子分类下'}), 400

                category.parent_id = new_parent_id

        db.session.commit()
        update_category_article_count(category.id)

        return jsonify({
            'success': True,
            'message': '分类更新成功'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"更新分类错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/categories/<int:cat_id>/archive', methods=['POST'])
@jwt_required()
def archive_category(cat_id):
    """归档分类"""
    try:
        category = KnowledgeCategory.query.get(cat_id)
        if not category:
            return jsonify({'success': False, 'error': '分类不存在'}), 404

        if not can_edit_category(category):
            return jsonify({'success': False, 'error': '无权限操作此分类'}), 403

        category.is_archived = 1 if not category.is_archived else 0
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '分类归档成功' if category.is_archived else '分类已恢复'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"归档分类错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/categories/<int:cat_id>', methods=['DELETE'])
@jwt_required()
def delete_category(cat_id):
    """删除分类"""
    try:
        category = KnowledgeCategory.query.get(cat_id)
        if not category:
            return jsonify({'success': False, 'error': '分类不存在'}), 404

        if not can_edit_category(category):
            return jsonify({'success': False, 'error': '无权限删除此分类'}), 403

        data = request.get_json() or {}
        action = data.get('action', 'delete')

        child_ids = get_all_child_category_ids(cat_id)
        articles = KnowledgeArticle.query.filter(KnowledgeArticle.category_id.in_(child_ids)).all()

        if articles:
            if action == 'migrate' and data.get('target_category_id'):
                target_category = KnowledgeCategory.query.get(data['target_category_id'])
                if not target_category:
                    return jsonify({'success': False, 'error': '目标分类不存在'}), 400
                for article in articles:
                    article.category_id = data['target_category_id']
            else:
                for article in articles:
                    article.category_id = None

        for child in category.children:
            child.parent_id = category.parent_id

        db.session.delete(category)
        db.session.commit()

        if category.parent_id:
            update_category_article_count(category.parent_id)

        return jsonify({
            'success': True,
            'message': '分类删除成功'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"删除分类错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/categories/<int:cat_id>/articles', methods=['GET'])
@jwt_required()
def get_category_articles(cat_id):
    """获取分类下的所有文章（包含子分类）"""
    try:
        category = KnowledgeCategory.query.get(cat_id)
        if not category:
            return jsonify({'success': False, 'error': '分类不存在'}), 404

        child_ids = get_all_child_category_ids(cat_id)
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        sort_by = request.args.get('sort_by', 'newest')
        status = request.args.get('status', 'published')

        query = KnowledgeArticle.query.filter(KnowledgeArticle.category_id.in_(child_ids))

        if status != 'all':
            query = query.filter_by(status=status)

        if sort_by == 'newest':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.updated_at.desc())
        elif sort_by == 'oldest':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.created_at.asc())
        elif sort_by == 'title':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.title)
        elif sort_by == 'views':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.view_count.desc())
        else:
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.updated_at.desc())

        total = query.count()
        articles = query.offset((page - 1) * per_page).limit(per_page).all()
        
        from sqlalchemy import func
        comment_counts = db.session.query(
            KnowledgeComment.article_id, func.count(KnowledgeComment.id).label('count')
        ).group_by(KnowledgeComment.article_id).all()
        comment_count_map = {count.article_id: count.count for count in comment_counts}

        return jsonify({
            'success': True,
            'data': {
                'articles': [art.to_dict(comment_count=comment_count_map.get(art.id, 0)) for art in articles],
                'total': total,
                'page': page,
                'per_page': per_page,
                'category': category.to_dict()
            }
        })
    except Exception as e:
        logger.error(f"获取分类文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 文章管理 ====================
@knowledge_bp.route('/articles', methods=['GET'])
@jwt_required()
def get_articles():
    """获取文章列表"""
    try:
        category_id = request.args.get('category_id', type=int)
        keyword = request.args.get('keyword')
        tag = request.args.get('tag')
        status = request.args.get('status')
        sort_by = request.args.get('sort_by', 'newest')
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        favorites = request.args.get('favorites', 'false').lower() == 'true'
        is_pinned = request.args.get('is_pinned', 'false').lower() == 'true'

        query = KnowledgeArticle.query

        if category_id:
            child_ids = get_all_child_category_ids(category_id)
            query = query.filter(KnowledgeArticle.category_id.in_(child_ids))
        if keyword:
            query = query.filter(
                or_(
                    KnowledgeArticle.title.ilike(f'%{keyword}%'),
                    KnowledgeArticle.content.ilike(f'%{keyword}%'),
                    KnowledgeArticle.summary.ilike(f'%{keyword}%')
                )
            )
        if tag:
            query = query.filter(KnowledgeArticle.tags.ilike(f'%"{tag}"%'))
        if status:
            query = query.filter_by(status=status)
        
        # 筛选置顶文章
        if is_pinned:
            query = query.filter_by(is_pinned=True)
        
        # 筛选收藏的文章
        if favorites:
            current_user_id = get_jwt_identity()
            logger.info(f"筛选收藏文章，用户 ID: {current_user_id}")
            favorite_article_ids = db.session.query(KnowledgeFavorite.article_id).filter_by(user_id=current_user_id).all()
            favorite_article_ids = [item[0] for item in favorite_article_ids]
            logger.info(f"用户收藏的文章 ID 列表：{favorite_article_ids}")
            if favorite_article_ids:
                query = query.filter(KnowledgeArticle.id.in_(favorite_article_ids))
            else:
                # 如果没有收藏任何文章，返回空结果
                logger.info("用户没有收藏任何文章，返回空结果")
                return jsonify({
                    'success': True,
                    'data': {
                        'articles': [],
                        'total': 0,
                        'page': page,
                        'per_page': per_page
                    }
                })

        if sort_by == 'newest':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.updated_at.desc())
        elif sort_by == 'oldest':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.created_at.asc())
        elif sort_by == 'title':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.title)
        elif sort_by == 'views':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.view_count.desc())
        else:
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.updated_at.desc())

        total = query.count()
        articles = query.offset((page - 1) * per_page).limit(per_page).all()
        
        from sqlalchemy import func
        comment_counts = db.session.query(
            KnowledgeComment.article_id, func.count(KnowledgeComment.id).label('count')
        ).group_by(KnowledgeComment.article_id).all()
        comment_count_map = {count.article_id: count.count for count in comment_counts}
        
        current_user_id = get_jwt_identity()
        favorite_article_ids = set()
        if current_user_id:
            favorites = KnowledgeFavorite.query.filter_by(user_id=current_user_id).all()
            favorite_article_ids = {f.article_id for f in favorites}
        
        articles_data = []
        for art in articles:
            art_dict = art.to_dict(comment_count=comment_count_map.get(art.id, 0))
            art_dict['is_favorited'] = art.id in favorite_article_ids
            articles_data.append(art_dict)

        return jsonify({
            'success': True,
            'data': {
                'articles': articles_data,
                'total': total,
                'page': page,
                'per_page': per_page
            }
        })
    except Exception as e:
        logger.error(f"获取文章列表错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/my', methods=['GET'])
@jwt_required()
def get_my_articles():
    """获取当前用户的文章"""
    try:
        current_user_id = get_jwt_identity()
        sort_by = request.args.get('sort_by', 'newest')
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        status = request.args.get('status')

        query = KnowledgeArticle.query.filter_by(author_id=current_user_id)

        if status:
            query = query.filter_by(status=status)

        if sort_by == 'newest':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.updated_at.desc())
        elif sort_by == 'oldest':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.created_at.asc())
        elif sort_by == 'title':
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.title)
        else:
            query = query.order_by(KnowledgeArticle.is_pinned.desc(), KnowledgeArticle.updated_at.desc())

        total = query.count()
        articles = query.offset((page - 1) * per_page).limit(per_page).all()
        
        from sqlalchemy import func
        comment_counts = db.session.query(
            KnowledgeComment.article_id, func.count(KnowledgeComment.id).label('count')
        ).group_by(KnowledgeComment.article_id).all()
        comment_count_map = {count.article_id: count.count for count in comment_counts}

        return jsonify({
            'success': True,
            'data': {
                'articles': [art.to_dict(comment_count=comment_count_map.get(art.id, 0)) for art in articles],
                'total': total,
                'page': page,
                'per_page': per_page
            }
        })
    except Exception as e:
        logger.error(f"获取我的文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles', methods=['POST'])
@jwt_required()
def create_article():
    """创建文章"""
    try:
        data = request.get_json()
        current_user_id = get_jwt_identity()
        user = get_current_user()

        if not data.get('title', '').strip():
            return jsonify({'success': False, 'error': '文章标题不能为空'}), 400
        if not data.get('content', '').strip():
            return jsonify({'success': False, 'error': '文章内容不能为空'}), 400

        category_id = data.get('category_id')
        if category_id:
            category = KnowledgeCategory.query.get(category_id)
            if not category:
                return jsonify({'success': False, 'error': '分类不存在'}), 400
            if category.children:
                return jsonify({'success': False, 'error': '文章只能属于叶节点分类'}), 400

        request_author_id = data.get('author_id')
        if request_author_id is not None:
            try:
                request_author_id = int(request_author_id)
            except (ValueError, TypeError):
                request_author_id = None
        
        if request_author_id is not None and check_admin():
            request_author = User.query.get(request_author_id)
            if request_author:
                final_author_id = request_author_id
                final_author_name = request_author.username
            else:
                final_author_id = current_user_id
                final_author_name = user.username if user else '未知'
        elif request_author_id is None and check_admin():
            final_author_id = None
            final_author_name = None
        else:
            final_author_id = current_user_id
            final_author_name = user.username if user else '未知'

        article = KnowledgeArticle(
            title=data['title'].strip(),
            content=data['content'],
            summary=data.get('summary', data['content'][:200] + '...' if len(data['content']) > 200 else data['content']),
            cover_image=data.get('cover_image'),
            category_id=category_id,
            tags=json.dumps(data.get('tags', []), ensure_ascii=False),
            author_id=final_author_id,
            author_name=final_author_name,
            status=data.get('status', 'draft')
        )
        db.session.add(article)
        db.session.commit()

        if article.status == 'published' and article.category_id:
            update_category_article_count(article.category_id)

        return jsonify({
            'success': True,
            'message': '文章创建成功',
            'data': {'id': article.id}
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"创建文章错误：{str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 文章详细操作路由（必须在通用路由之前） ====================
@knowledge_bp.route('/articles/<int:art_id>/comments', methods=['GET'])
@jwt_required()
def get_article_comments(art_id):
    """获取文章评论列表"""
    import traceback
    try:
        logger.info(f"开始获取文章 {art_id} 的评论")
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)

        query = KnowledgeComment.query.options(
            joinedload(KnowledgeComment.user)
        ).filter_by(
            article_id=art_id,
            parent_id=None
        ).order_by(KnowledgeComment.created_at.desc())
        
        total = query.count()
        comments_query = query.offset((page - 1) * per_page).limit(per_page)

        comments = []
        for comment in comments_query:
            replies_list = KnowledgeComment.query.options(
                joinedload(KnowledgeComment.user)
            ).filter_by(
                parent_id=comment.id
            ).order_by(KnowledgeComment.created_at.asc()).all()
            
            comment_dict = {
                'id': comment.id,
                'article_id': comment.article_id,
                'user_id': comment.user_id,
                'user_name': comment.user.username if comment and comment.user else '未知',
                'user_avatar': comment.user.avatar if comment and comment.user else None,
                'content': comment.content,
                'block_id': comment.block_id,
                'parent_id': comment.parent_id,
                'is_resolved': bool(comment.is_resolved),
                'created_at': comment.created_at.isoformat() if comment.created_at else None,
                'updated_at': comment.updated_at.isoformat() if comment.updated_at else None,
                'replies': []
            }
            for reply in replies_list:
                comment_dict['replies'].append({
                    'id': reply.id,
                    'article_id': reply.article_id,
                    'user_id': reply.user_id,
                    'user_name': reply.user.username if reply and reply.user else '未知',
                    'user_avatar': reply.user.avatar if reply and reply.user else None,
                    'content': reply.content,
                    'block_id': reply.block_id,
                    'parent_id': reply.parent_id,
                    'is_resolved': bool(reply.is_resolved),
                    'created_at': reply.created_at.isoformat() if reply.created_at else None,
                    'updated_at': reply.updated_at.isoformat() if reply.updated_at else None
                })
            comments.append(comment_dict)

        total_pages = (total + per_page - 1) // per_page if per_page > 0 else 1
        logger.info(f"成功获取文章 {art_id} 的评论，共 {len(comments)} 条")
        return jsonify({
            'success': True,
            'data': {
                'comments': comments,
                'total': total,
                'pages': total_pages,
                'current_page': page
            }
        })
    except Exception as e:
        logger.error(f"获取评论错误：{str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/comments', methods=['POST'])
@jwt_required()
def create_article_comment(art_id):
    """发表评论"""
    import traceback
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        current_user_id = get_jwt_identity()
        data = request.get_json()

        content = data.get('content', '').strip()
        if not content:
            return jsonify({'success': False, 'error': '评论内容不能为空'}), 400

        block_id = data.get('block_id')
        parent_id = data.get('parent_id')

        if parent_id:
            parent_comment = KnowledgeComment.query.get(parent_id)
            if not parent_comment or parent_comment.article_id != art_id:
                return jsonify({'success': False, 'error': '父评论不存在'}), 400

        comment = KnowledgeComment(
            article_id=art_id,
            user_id=current_user_id,
            content=content,
            block_id=block_id,
            parent_id=parent_id
        )
        db.session.add(comment)
        db.session.commit()
        
        # Refresh to get the user relationship
        db.session.refresh(comment)

        return jsonify({
            'success': True,
            'message': '评论成功',
            'data': {
                'id': comment.id,
                'article_id': comment.article_id,
                'user_id': comment.user_id,
                'user_name': comment.user.username if comment and comment.user else '未知',
                'user_avatar': comment.user.avatar if comment and comment.user else None,
                'content': comment.content,
                'block_id': comment.block_id,
                'parent_id': comment.parent_id,
                'is_resolved': bool(comment.is_resolved),
                'created_at': comment.created_at.isoformat() if comment.created_at else None
            }
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"发表评论错误：{str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/comments/<int:comment_id>', methods=['DELETE'])
@jwt_required()
def delete_article_comment(art_id, comment_id):
    """删除评论"""
    try:
        current_user_id = get_jwt_identity()
        is_admin = check_admin()

        comment = KnowledgeComment.query.filter_by(id=comment_id, article_id=art_id).first()
        if not comment:
            return jsonify({'success': False, 'error': '评论不存在'}), 404

        if comment.user_id != current_user_id and not is_admin:
            return jsonify({'success': False, 'error': '无权限删除此评论'}), 403

        if comment.replies:
            comment.content = '[已删除]'
            comment.is_resolved = True
        else:
            db.session.delete(comment)

        db.session.commit()

        return jsonify({
            'success': True,
            'message': '删除成功'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"删除评论错误：{str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/related', methods=['GET'])
@jwt_required()
def get_related_articles(art_id):
    """获取相关文章"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        limit = request.args.get('limit', 5, type=int)

        query = KnowledgeArticle.query.filter(
            KnowledgeArticle.id != art_id,
            KnowledgeArticle.status == 'published'
        )

        if article.category_id:
            query = query.filter_by(category_id=article.category_id)

        articles = query.order_by(KnowledgeArticle.view_count.desc()).limit(limit).all()
        
        from sqlalchemy import func
        comment_counts = db.session.query(
            KnowledgeComment.article_id, func.count(KnowledgeComment.id).label('count')
        ).group_by(KnowledgeComment.article_id).all()
        comment_count_map = {count.article_id: count.count for count in comment_counts}

        return jsonify({
            'success': True,
            'data': {
                'items': [art.to_dict(comment_count=comment_count_map.get(art.id, 0)) for art in articles]
            }
        })
    except Exception as e:
        logger.error(f"获取相关文章错误：{str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/export/<export_type>', methods=['GET'])
@jwt_required()
def export_article(art_id, export_type):
    """导出文章为 PDF 或 DOCX"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        if export_type not in ['pdf', 'docx']:
            return jsonify({'success': False, 'error': '不支持的导出格式'}), 400

        from flask import make_response
        import io

        if export_type == 'pdf':
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
                from reportlab.lib.units import cm
                from reportlab.lib.enums import TA_LEFT
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from io import BytesIO
                import re

                buffer = BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm)
                styles = getSampleStyleSheet()
                story = []

                # 尝试注册中文字体
                try:
                    # 尝试使用系统字体
                    import platform
                    system_name = platform.system()
                    if system_name == 'Windows':
                        font_path = r'C:\Windows\Fonts\simhei.ttf'
                        if os.path.exists(font_path):
                            pdfmetrics.registerFont(TTFont('Chinese', font_path))
                            chinese_font = 'Chinese'
                        else:
                            chinese_font = 'Helvetica'
                    elif system_name == 'Darwin':
                        font_path = '/System/Library/Fonts/STHeiti Light.ttc'
                        if os.path.exists(font_path):
                            pdfmetrics.registerFont(TTFont('Chinese', font_path))
                            chinese_font = 'Chinese'
                        else:
                            chinese_font = 'Helvetica'
                    else:
                        chinese_font = 'Helvetica'
                except Exception as font_error:
                    logger.warning(f"字体加载失败：{font_error}")
                    chinese_font = 'Helvetica'

                # 标题样式
                title_style = styles['Title']
                title_style.fontSize = 18
                title_style.fontName = chinese_font
                title_style.alignment = TA_LEFT
                story.append(Paragraph(article.title, title_style))
                story.append(Spacer(1, 0.5*cm))

                # 元信息样式
                meta_style = styles['Normal']
                meta_style.fontSize = 10
                meta_style.fontName = chinese_font
                try:
                    category_name = article.category.name if article.category else '未分类'
                except Exception:
                    category_name = '未分类'
                meta = f"作者：{article.author_name or '未知'} | 分类：{category_name}"
                story.append(Paragraph(meta, meta_style))
                story.append(Spacer(1, 0.5*cm))
                story.append(PageBreak())

                # 内容样式
                content_style = styles['Normal']
                content_style.fontSize = 12
                content_style.fontName = chinese_font

                # 处理文章内容 - 清理 HTML 标签
                content_text = article.content or ''
                # 移除 HTML 标签
                clean_content = re.sub(r'<[^>]+>', '', content_text)
                # 替换常见的 HTML 实体
                clean_content = clean_content.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
                
                # 分段处理内容
                paragraphs = clean_content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), content_style))
                        story.append(Spacer(1, 0.2*cm))

                doc.build(story)
                buffer.seek(0)

                response = make_response(buffer.read())
                response.headers['Content-Type'] = 'application/pdf'
                from urllib.parse import quote
                safe_filename = quote(article.title, safe='')
                response.headers['Content-Disposition'] = f'attachment; filename={safe_filename}.pdf'
                return response
            except ImportError as ie:
                logger.error(f"PDF 导出库未安装：{ie}")
                return jsonify({'success': False, 'error': 'PDF 导出功能未安装 reportlab 库'}), 500
            except Exception as pdf_error:
                logger.error(f"PDF 生成失败：{pdf_error}")
                return jsonify({'success': False, 'error': f'PDF 生成失败：{str(pdf_error)}'}), 500

        elif export_type == 'docx':
            from io import BytesIO
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                doc = Document()
                doc.add_heading(article.title, 0)

                p = doc.add_paragraph()
                p.add_run(f'作者：{article.author_name}').bold = True
                try:
                    category_name = article.category.name if article.category else '未分类'
                except Exception:
                    category_name = '未分类'
                p.add_run(f' | 分类：{category_name}')

                doc.add_paragraph(article.content)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                response = make_response(buffer.read())
                response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                from urllib.parse import quote
                safe_filename = quote(article.title, safe='')
                response.headers['Content-Disposition'] = f'attachment; filename={safe_filename}.docx'
                return response
            except ImportError:
                return jsonify({'success': False, 'error': '导出功能未安装 python-docx 库'}), 500

    except Exception as e:
        logger.error(f"导出文章错误：{str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/attachments', methods=['POST'])
@jwt_required()
def upload_attachment(art_id):
    """上传附件"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        if not can_edit_article(article):
            return jsonify({'success': False, 'error': '无权限上传附件到此文章'}), 403

        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '没有上传文件'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '文件名为空'}), 400

        upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads', 'knowledge')
        os.makedirs(upload_dir, exist_ok=True)

        file_ext = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4().hex}{file_ext}"
        file_path = os.path.join(upload_dir, unique_filename)
        file.save(file_path)

        attachment = KnowledgeAttachment(
            article_id=art_id,
            filename=file.filename,
            storage_name=unique_filename,
            file_path=file_path,
            file_size=os.path.getsize(file_path)
        )
        db.session.add(attachment)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '文件上传成功',
            'data': {
                'id': attachment.id,
                'filename': attachment.filename,
                'file_size': attachment.file_size
            }
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"上传附件错误：{str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/attachments/<int:att_id>', methods=['GET'])
@jwt_required()
def download_attachment(art_id, att_id):
    """下载附件"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        attachment = KnowledgeAttachment.query.filter_by(id=att_id, article_id=art_id).first()
        if not attachment:
            return jsonify({'success': False, 'error': '附件不存在'}), 404

        if not os.path.exists(attachment.file_path):
            return jsonify({'success': False, 'error': '文件不存在'}), 404

        return send_file(attachment.file_path, as_attachment=True, download_name=attachment.filename)
    except Exception as e:
        logger.error(f"下载附件错误：{str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/attachments/<int:att_id>', methods=['DELETE'])
@jwt_required()
def delete_attachment(art_id, att_id):
    """删除附件"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        if not can_edit_article(article):
            return jsonify({'success': False, 'error': '无权限删除此附件'}), 403

        attachment = KnowledgeAttachment.query.filter_by(id=att_id, article_id=art_id).first()
        if not attachment:
            return jsonify({'success': False, 'error': '附件不存在'}), 404

        if os.path.exists(attachment.file_path):
            try:
                os.remove(attachment.file_path)
            except:
                pass

        db.session.delete(attachment)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '附件删除成功'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"删除附件错误：{str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 文章通用路由（必须在最后） ====================
@knowledge_bp.route('/articles/<int:art_id>', methods=['GET'])
@jwt_required()
def get_article(art_id):
    """获取文章详情"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        article.view_count = (article.view_count or 0) + 1
        db.session.commit()
        
        current_user_id = get_jwt_identity()
        is_favorited = KnowledgeFavorite.query.filter_by(
            article_id=art_id,
            user_id=current_user_id
        ).first() is not None
        
        result = article.to_dict(include_content=True)
        result['is_favorited'] = is_favorited

        return jsonify({
            'success': True,
            'data': result
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"获取文章详情错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>', methods=['PUT'])
@jwt_required()
def update_article(art_id):
    """更新文章"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        if not can_edit_article(article):
            return jsonify({'success': False, 'error': '无权限编辑此文章'}), 403

        data = request.get_json()
        old_category_id = article.category_id
        old_status = article.status
        new_category_id = data.get('category_id')
        new_status = data.get('status', old_status)

        if new_category_id and new_category_id != old_category_id:
            new_category = KnowledgeCategory.query.get(new_category_id)
            if not new_category:
                return jsonify({'success': False, 'error': '分类不存在'}), 400
            if new_category.children:
                return jsonify({'success': False, 'error': '文章只能属于叶节点分类'}), 400

        if 'title' in data:
            if not data['title'].strip():
                return jsonify({'success': False, 'error': '文章标题不能为空'}), 400
            article.title = data['title'].strip()
        if 'content' in data:
            if not data['content'].strip():
                return jsonify({'success': False, 'error': '文章内容不能为空'}), 400
            article.content = data['content']
        if 'summary' in data:
            article.summary = data['summary']
        if 'cover_image' in data:
            article.cover_image = data['cover_image']
        if 'category_id' in data:
            article.category_id = new_category_id
        if 'tags' in data:
            article.tags = json.dumps(data['tags'], ensure_ascii=False)
        if 'status' in data:
            article.status = data['status']
        if 'is_public' in data:
            article.is_public = 1 if data['is_public'] else 0
        
        # 处理作者信息更新
        new_author_id = data.get('author_id')
        logger.info(f"update_article: raw author_id from data = {data.get('author_id')}, parsed = {new_author_id}, is_admin = {check_admin()}")
        logger.info(f"update_article: original author_id={article.author_id}, author_name={article.author_name}")
        
        if new_author_id is not None:
            try:
                new_author_id = int(new_author_id)
            except (ValueError, TypeError):
                new_author_id = None
            
            if new_author_id is not None and check_admin():
                new_author = User.query.get(new_author_id)
                logger.info(f"update_article: new_author query result = {new_author}")
                if new_author:
                    article.author_id = new_author_id
                    article.author_name = new_author.username
                    logger.info(f"update_article: updated to author_id={article.author_id}, author_name={article.author_name}")
                else:
                    logger.warning(f"update_article: user {new_author_id} not found, keeping original author")
            elif new_author_id is None and check_admin():
                article.author_id = None
                article.author_name = None
                logger.info(f"update_article: cleared author info")
        elif check_admin():
            # 管理员编辑时，如果没有传递 author_id，保持原有作者信息不变
            logger.info(f"update_article: admin editing without author_id, keeping original author")
            pass
        else:
            # 非管理员编辑时，强制使用当前用户作为作者
            current_user_id = get_jwt_identity()
            current_user = get_current_user()
            article.author_id = current_user_id
            article.author_name = current_user.username if current_user else '未知'
            logger.info(f"update_article: non-admin user, forced author_id={article.author_id}, author_name={article.author_name}")

        db.session.commit()
        logger.info(f"update_article: after commit - author_id={article.author_id}, author_name={article.author_name}")

        if old_category_id:
            update_category_article_count(old_category_id)
        if new_category_id and new_category_id != old_category_id:
            update_category_article_count(new_category_id)
        elif new_status != old_status:
            if new_category_id:
                update_category_article_count(new_category_id)

        return jsonify({
            'success': True,
            'message': '文章更新成功'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"更新文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/status', methods=['PUT'])
@jwt_required()
def update_article_status(art_id):
    """更新文章状态"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        if not can_edit_article(article):
            return jsonify({'success': False, 'error': '无权限操作此文章'}), 403

        data = request.get_json()
        new_status = data.get('status')
        if new_status not in ['draft', 'published']:
            return jsonify({'success': False, 'error': '无效的状态'}), 400

        old_status = article.status
        article.status = new_status
        db.session.commit()

        if article.category_id:
            if old_status != new_status:
                update_category_article_count(article.category_id)

        return jsonify({
            'success': True,
            'message': '状态更新成功'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"更新文章状态错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/pin', methods=['POST'])
@jwt_required()
def toggle_article_pin(art_id):
    """切换文章置顶状态"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        if not can_edit_article(article):
            return jsonify({'success': False, 'error': '无权限操作此文章'}), 403

        data = request.get_json()
        is_pinned = data.get('is_pinned', False)
        
        article.is_pinned = 1 if is_pinned else 0
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '置顶成功' if is_pinned else '取消置顶成功',
            'data': {
                'is_pinned': bool(article.is_pinned)
            }
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"置顶操作错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>', methods=['DELETE'])
@jwt_required()
def delete_article(art_id):
    """删除文章"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        if not can_edit_article(article):
            return jsonify({'success': False, 'error': '无权限删除此文章'}), 403

        old_category_id = article.category_id

        attachments = KnowledgeAttachment.query.filter_by(article_id=art_id).all()
        for att in attachments:
            if os.path.exists(att.file_path):
                try:
                    os.remove(att.file_path)
                except:
                    pass
            db.session.delete(att)

        comments = KnowledgeComment.query.filter_by(article_id=art_id).all()
        for comment in comments:
            db.session.delete(comment)

        db.session.delete(article)
        db.session.commit()

        if old_category_id:
            update_category_article_count(old_category_id)

        return jsonify({
            'success': True,
            'message': '文章删除成功'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"删除文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 批量操作 ====================
@knowledge_bp.route('/articles/batch/move', methods=['POST'])
@jwt_required()
def batch_move_articles():
    """批量移动文章"""
    try:
        data = request.get_json()
        article_ids = data.get('article_ids', [])
        target_category_id = data.get('target_category_id')

        if not article_ids:
            return jsonify({'success': False, 'error': '请选择要移动的文章'}), 400

        if target_category_id:
            target_category = KnowledgeCategory.query.get(target_category_id)
            if not target_category:
                return jsonify({'success': False, 'error': '目标分类不存在'}), 400
            if target_category.children:
                return jsonify({'success': False, 'error': '文章只能移动到叶节点分类'}), 400

        affected_categories = set()
        for art_id in article_ids:
            article = KnowledgeArticle.query.get(art_id)
            if article:
                if not can_edit_article(article):
                    return jsonify({'success': False, 'error': f'无权限移动文章{art_id}'}), 403
                if article.category_id:
                    affected_categories.add(article.category_id)
                article.category_id = target_category_id
                affected_categories.add(target_category_id) if target_category_id else None

        db.session.commit()

        for cat_id in affected_categories:
            if cat_id:
                update_category_article_count(cat_id)

        return jsonify({
            'success': True,
            'message': f'成功移动{len(article_ids)}篇文章'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"批量移动文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/batch/status', methods=['POST'])
@jwt_required()
def batch_update_status():
    """批量更新文章状态"""
    try:
        data = request.get_json()
        article_ids = data.get('article_ids', [])
        new_status = data.get('status')

        if not article_ids:
            return jsonify({'success': False, 'error': '请选择要更新的文章'}), 400
        if new_status not in ['draft', 'published']:
            return jsonify({'success': False, 'error': '无效的状态'}), 400

        affected_categories = set()
        for art_id in article_ids:
            article = KnowledgeArticle.query.get(art_id)
            if article:
                if not can_edit_article(article):
                    return jsonify({'success': False, 'error': f'无权限更新文章{art_id}'}), 403
                if article.category_id and article.status != new_status:
                    affected_categories.add(article.category_id)
                article.status = new_status

        db.session.commit()

        for cat_id in affected_categories:
            if cat_id:
                update_category_article_count(cat_id)

        return jsonify({
            'success': True,
            'message': f'成功更新{len(article_ids)}篇文章状态'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"批量更新状态错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/batch/delete', methods=['POST'])
@jwt_required()
def batch_delete_articles():
    """批量删除文章"""
    try:
        data = request.get_json()
        article_ids = data.get('article_ids', [])

        if not article_ids:
            return jsonify({'success': False, 'error': '请选择要删除的文章'}), 400

        affected_categories = set()
        deleted_count = 0
        for art_id in article_ids:
            article = KnowledgeArticle.query.get(art_id)
            if article:
                if not can_edit_article(article):
                    return jsonify({'success': False, 'error': f'无权限删除文章{art_id}'}), 403
                if article.category_id:
                    affected_categories.add(article.category_id)

                attachments = KnowledgeAttachment.query.filter_by(article_id=art_id).all()
                for att in attachments:
                    if os.path.exists(att.file_path):
                        try:
                            os.remove(att.file_path)
                        except:
                            pass
                    db.session.delete(att)

                comments = KnowledgeComment.query.filter_by(article_id=art_id).all()
                for comment in comments:
                    db.session.delete(comment)

                db.session.delete(article)
                deleted_count += 1

        db.session.commit()

        for cat_id in affected_categories:
            if cat_id:
                update_category_article_count(cat_id)

        return jsonify({
            'success': True,
            'message': f'成功删除{deleted_count}篇文章'
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"批量删除文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/like', methods=['POST'])
@jwt_required()
def like_article(art_id):
    """点赞文章"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        article.like_count = (article.like_count or 0) + 1
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '点赞成功',
            'data': {'like_count': article.like_count}
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"点赞文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/favorite', methods=['POST'])
@jwt_required()
def favorite_article(art_id):
    """收藏文章"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        current_user_id = get_jwt_identity()
        
        existing = KnowledgeFavorite.query.filter_by(
            article_id=art_id,
            user_id=current_user_id
        ).first()
        
        if existing:
            return jsonify({'success': False, 'error': '已经收藏过了'}), 400
        
        favorite = KnowledgeFavorite(
            article_id=art_id,
            user_id=current_user_id
        )
        db.session.add(favorite)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '收藏成功',
            'data': {'is_favorited': True}
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"收藏文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@knowledge_bp.route('/articles/<int:art_id>/favorite', methods=['DELETE'])
@jwt_required()
def unfavorite_article(art_id):
    """取消收藏文章"""
    try:
        article = KnowledgeArticle.query.get(art_id)
        if not article:
            return jsonify({'success': False, 'error': '文章不存在'}), 404

        current_user_id = get_jwt_identity()
        
        favorite = KnowledgeFavorite.query.filter_by(
            article_id=art_id,
            user_id=current_user_id
        ).first()
        
        if not favorite:
            return jsonify({'success': False, 'error': '尚未收藏'}), 400
        
        db.session.delete(favorite)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': '取消收藏成功',
            'data': {'is_favorited': False}
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"取消收藏错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 全文搜索 ====================
@knowledge_bp.route('/search', methods=['GET'])
@jwt_required()
def search_articles():
    """全文搜索"""
    try:
        keyword = request.args.get('q', '')
        category_id = request.args.get('category_id', type=int)

        if not keyword:
            return jsonify({'success': False, 'error': '请输入搜索关键词'}), 400

        query = KnowledgeArticle.query.filter(
            or_(
                KnowledgeArticle.title.ilike(f'%{keyword}%'),
                KnowledgeArticle.content.ilike(f'%{keyword}%'),
                KnowledgeArticle.summary.ilike(f'%{keyword}%'),
                KnowledgeArticle.tags.ilike(f'%{keyword}%')
            )
        ).filter_by(status='published')

        if category_id:
            child_ids = get_all_child_category_ids(category_id)
            query = query.filter(KnowledgeArticle.category_id.in_(child_ids))

        articles = query.all()
        results = []
        keyword_lower = keyword.lower()
        
        from sqlalchemy import func
        comment_counts = db.session.query(
            KnowledgeComment.article_id, func.count(KnowledgeComment.id).label('count')
        ).group_by(KnowledgeComment.article_id).all()
        comment_count_map = {count.article_id: count.count for count in comment_counts}

        for art in articles:
            score = 0
            title_match = keyword_lower in art.title.lower()

            if title_match:
                score += 10
            if keyword_lower in (art.content or '').lower():
                score += 5
            if keyword_lower in (art.summary or '').lower():
                score += 3

            try:
                tags_list = json.loads(art.tags) if art.tags else []
            except json.JSONDecodeError:
                tags_list = []
            tag_match = any(keyword_lower in tag.lower() for tag in tags_list)
            if tag_match:
                score += 2

            content = art.content or ''
            snippet = content
            if keyword_lower in content.lower():
                idx = content.lower().find(keyword_lower)
                start = max(0, idx - 50)
                end = min(len(content), idx + len(keyword) + 50)
                snippet = '...' + content[start:end] + '...'

            highlighted_title = art.title.replace(keyword, f'<em>{keyword}</em>')

            results.append({
                'id': art.id,
                'title': art.title,
                'highlighted_title': highlighted_title,
                'snippet': snippet[:200],
                'category_id': art.category_id,
                'category_name': art.category.name if art.category else '未分类',
                'category_path': art.category_path if hasattr(art, 'category_path') else art.category.name if art.category else '未分类',
                'score': score,
                'author_name': art.author_name,
                'created_at': art.created_at.isoformat() if art.created_at else None,
                'view_count': art.view_count,
                'comment_count': comment_count_map.get(art.id, 0)
            })

        results.sort(key=lambda x: x['score'], reverse=True)

        return jsonify({
            'success': True,
            'data': {
                'keyword': keyword,
                'total': len(results),
                'results': results[:20]
            }
        })
    except Exception as e:
        logger.error(f"搜索文章错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 标签管理 ====================
@knowledge_bp.route('/tags', methods=['GET'])
@jwt_required()
def get_tags():
    """获取所有标签"""
    try:
        articles = KnowledgeArticle.query.filter_by(status='published').all()
        tag_count = {}
        for art in articles:
            try:
                tags_list = json.loads(art.tags) if art.tags else []
            except json.JSONDecodeError:
                tags_list = []
            for tag in tags_list:
                tag_count[tag] = tag_count.get(tag, 0) + 1

        tags = [{'name': tag, 'count': count} for tag, count in tag_count.items()]
        tags.sort(key=lambda x: x['count'], reverse=True)

        return jsonify({
            'success': True,
            'data': tags
        })
    except Exception as e:
        logger.error(f"获取标签错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 统计 ====================
@knowledge_bp.route('/statistics', methods=['GET'])
@jwt_required()
def get_statistics():
    """获取知识库统计"""
    try:
        total_articles = KnowledgeArticle.query.filter_by(status='published').count()
        draft_articles = KnowledgeArticle.query.filter_by(status='draft').count()
        total_categories = KnowledgeCategory.query.filter_by(is_archived=0).count()

        articles = KnowledgeArticle.query.filter_by(status='published').all()
        total_views = sum(art.view_count or 0 for art in articles)
        total_likes = sum(art.like_count or 0 for art in articles)

        hot_articles = KnowledgeArticle.query.filter_by(status='published').order_by(KnowledgeArticle.view_count.desc()).limit(10).all()
        recent_articles = KnowledgeArticle.query.filter_by(status='published').order_by(KnowledgeArticle.updated_at.desc()).limit(10).all()

        return jsonify({
            'success': True,
            'data': {
                'total_articles': total_articles,
                'draft_articles': draft_articles,
                'total_categories': total_categories,
                'total_views': total_views,
                'total_likes': total_likes,
                'hot_articles': [{'id': art.id, 'title': art.title, 'view_count': art.view_count or 0} for art in hot_articles],
                'recent_articles': [{'id': art.id, 'title': art.title, 'updated_at': art.updated_at.isoformat() if art.updated_at else None} for art in recent_articles]
            }
        })
    except Exception as e:
        logger.error(f"获取统计错误: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500



@knowledge_bp.route('/stats', methods=['GET'])
@jwt_required()
def get_stats():
    """获取知识库统计"""
    try:
        current_user_id = get_jwt_identity()
        
        total = KnowledgeArticle.query.count()
        published = KnowledgeArticle.query.filter_by(status='published').count()
        draft = KnowledgeArticle.query.filter_by(status='draft').count()
        my_articles = KnowledgeArticle.query.filter_by(author_id=current_user_id).count()
        
        # 总浏览量
        total_views = db.session.query(func.sum(KnowledgeArticle.view_count)).scalar() or 0
        
        # 我的收藏统计
        favorites_count = KnowledgeFavorite.query.filter_by(user_id=current_user_id).count()
        
        return jsonify({
            'success': True,
            'data': {
                'total': total,
                'published': published,
                'draft': draft,
                'my': my_articles,
                'favorites': favorites_count,
                'views': total_views
            }
        })
    except Exception as e:
        logger.error(f"获取统计错误：{str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
